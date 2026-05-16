"""
src/utils/convert_doc.py
------------------------
Converts PROJECT_DOCUMENTATION.md into a professionally formatted Word document.

Features:
  - Title page with project branding
  - Color-coded heading hierarchy
  - Monospace code blocks with grey background
  - Formatted tables with header row styling
  - Blockquotes styled as tip/note/important callouts
  - Consistent body font, spacing, and margins
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ─── Color Palette ────────────────────────────────────────────────────────────
COLOR_TITLE      = RGBColor(0x1A, 0x23, 0x4E)   # Deep Navy Blue
COLOR_H1         = RGBColor(0x1A, 0x23, 0x4E)   # Navy
COLOR_H2         = RGBColor(0x16, 0x50, 0x8A)   # Rich Blue
COLOR_H3         = RGBColor(0x2E, 0x7D, 0x32)   # Forest Green
COLOR_H4         = RGBColor(0x6A, 0x1E, 0x55)   # Deep Purple
COLOR_CODE_BG    = RGBColor(0xF2, 0xF3, 0xF5)   # Light grey
COLOR_TIP_TEXT   = RGBColor(0x1B, 0x5E, 0x20)   # Dark Green
COLOR_NOTE_TEXT  = RGBColor(0x0D, 0x47, 0xA1)   # Dark Blue
COLOR_WARN_TEXT  = RGBColor(0x7F, 0x36, 0x00)   # Dark Orange
COLOR_TABLE_HDR  = RGBColor(0x1A, 0x23, 0x4E)   # Navy (matches H1)
BODY_FONT        = "Calibri"
CODE_FONT        = "Consolas"


def set_paragraph_shading(paragraph, fill_hex: str):
    """Set background color of a paragraph (for code blocks)."""
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    pPr.append(shd)


def set_cell_bg(cell, fill_hex: str):
    """Set background color for a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def add_horizontal_rule(doc):
    """Add a thin horizontal line to separate sections."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "AAAAAA")
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(6)


def add_title_page(doc):
    """Create a centered, styled title page."""
    # Spacer
    for _ in range(4):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)

    # Badge line
    badge = doc.add_paragraph()
    badge.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = badge.add_run("[FINAL YEAR PROJECT]")
    run.font.name = BODY_FONT
    run.font.size = Pt(13)
    run.font.color.rgb = COLOR_H2
    run.bold = True
    badge.paragraph_format.space_after = Pt(18)

    # Main title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("FEDERATED SELF-SUPERVISED LEARNING\n(FedSSL) FOR TB DETECTION")
    run.font.name = BODY_FONT
    run.font.size = Pt(28)
    run.font.color.rgb = COLOR_TITLE
    run.bold = True
    title.paragraph_format.space_after = Pt(20)

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("THE ABSOLUTE MASTER GUIDE")
    run.font.name = BODY_FONT
    run.font.size = Pt(16)
    run.italic = True
    run.font.color.rgb = COLOR_H2
    subtitle.paragraph_format.space_after = Pt(8)

    # Tagline
    tag = doc.add_paragraph()
    tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tag.add_run("Complete Study Notes · From Zero to Expert")
    run.font.name = BODY_FONT
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    tag.paragraph_format.space_after = Pt(40)

    # Divider
    add_horizontal_rule(doc)

    # Stats block
    stats = doc.add_paragraph()
    stats.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = stats.add_run(
        "Training Data: 46,718 NIH X-rays    |    "
        "Round 3 AUC: 0.8942    |    "
        "Hospitals Simulated: 5\n"
        "Backbone: ResNet50    |    Device: NVIDIA RTX 2050    |    Framework: PyTorch 2.5"
    )
    run.font.name = BODY_FONT
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    stats.paragraph_format.space_after = Pt(6)

    add_horizontal_rule(doc)
    doc.add_page_break()


def add_styled_heading(doc, text: str, level: int):
    """Add a heading with color, font size, and spacing."""
    p = doc.add_paragraph()

    config = {
        0: (COLOR_H1, 22, True, Pt(18), Pt(6)),    # H1 (# )
        1: (COLOR_H1, 18, True, Pt(14), Pt(4)),    # H2 (## )
        2: (COLOR_H2, 14, True, Pt(10), Pt(3)),    # H3 (### )
        3: (COLOR_H3, 12, True, Pt(8),  Pt(2)),    # H4 (#### )
        4: (COLOR_H4, 11, True, Pt(6),  Pt(2)),    # H5 (#####)
    }
    color, size, bold, space_before, space_after = config.get(level, config[4])

    run = p.add_run(text)
    run.font.name = BODY_FONT
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold

    p.paragraph_format.space_before = space_before
    p.paragraph_format.space_after = space_after

    # Underline for H1 only
    if level == 0:
        add_horizontal_rule(doc)

    return p


def add_code_block(doc, lines: list):
    """Add a shaded code block with monospace font."""
    # Top spacer
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(0)
    spacer.paragraph_format.space_before = Pt(3)

    for line in lines:
        p = doc.add_paragraph()
        set_paragraph_shading(p, "F2F3F5")   # Light grey hex
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

        # Strip markdown-style line
        display = line.rstrip("\n")
        run = p.add_run(display)
        run.font.name = CODE_FONT
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x24, 0x29, 0x3E)

    # Bottom spacer
    spacer2 = doc.add_paragraph()
    spacer2.paragraph_format.space_before = Pt(0)
    spacer2.paragraph_format.space_after = Pt(6)


def add_blockquote(doc, text: str, alert_type: str = ""):
    """Add styled callout block for TIP, NOTE, IMPORTANT, etc."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.35)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)

    color_map = {
        "TIP":       (COLOR_TIP_TEXT,  "TIP: "),
        "NOTE":      (COLOR_NOTE_TEXT, "NOTE: "),
        "IMPORTANT": (COLOR_WARN_TEXT, "IMPORTANT: "),
        "CAUTION":   (RGBColor(0xB7, 0x1C, 0x1C), "CAUTION: "),
        "WARNING":   (COLOR_WARN_TEXT, "WARNING: "),
    }

    color, label = color_map.get(alert_type.upper(), (COLOR_NOTE_TEXT, ""))

    if label:
        label_run = p.add_run(label)
        label_run.font.name = BODY_FONT
        label_run.font.size = Pt(10.5)
        label_run.font.color.rgb = color
        label_run.bold = True

    body_run = p.add_run(text)
    body_run.font.name = BODY_FONT
    body_run.font.size = Pt(10.5)
    body_run.font.color.rgb = color


def add_table(doc, rows: list):
    """Add a styled table with a navy header row."""
    if not rows:
        return
    num_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    for i, row_data in enumerate(rows):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = cell_text
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT

            for run in para.runs:
                run.font.name = BODY_FONT
                run.font.size = Pt(9.5)

            if i == 0:
                # Header row: navy background, white text, bold
                set_cell_bg(cell, "1A234E")
                for run in para.runs:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.bold = True
            elif i % 2 == 0:
                # Alternating row shading for readability
                set_cell_bg(cell, "EEF2F7")

    # Space after table
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)


def add_body_paragraph(doc, line: str):
    """Add a styled body paragraph with inline bold/italic/code support."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)

    # Check for bullet points
    stripped = line.strip()
    is_bullet = stripped.startswith("- ") or stripped.startswith("* ")
    if is_bullet:
        p.style = "List Bullet"
        line = stripped[2:]

    # Parse inline: bold (**), italic (*), inline code (`)
    pattern = re.compile(r"(\*\*.*?\*\*|\*.*?\*|`[^`]+`)")
    parts = pattern.split(line)

    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
            run.font.name = BODY_FONT
            run.font.size = Pt(10.5)
        elif part.startswith("`") and part.endswith("`"):
            run = p.add_run(part[1:-1])
            run.font.name = CODE_FONT
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)  # Inline code: raspberry red
        elif part.startswith("*") and part.endswith("*"):
            run = p.add_run(part[1:-1])
            run.italic = True
            run.font.name = BODY_FONT
            run.font.size = Pt(10.5)
        else:
            run = p.add_run(part)
            run.font.name = BODY_FONT
            run.font.size = Pt(10.5)


# ─── Main Parser ─────────────────────────────────────────────────────────────

def create_docx(md_path: str, docx_path: str):
    print(f"Reading {md_path}...")
    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    doc = Document()

    # Set page margins to 2.5 cm all around (A4-friendly)
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    # Apply default normal style
    normal_style = doc.styles["Normal"]
    normal_style.font.name = BODY_FONT
    normal_style.font.size = Pt(10.5)

    add_title_page(doc)

    print("Parsing and styling document...")

    in_code_block = False
    code_lines    = []
    in_table      = False
    table_rows    = []
    skip_next     = False   # Used to skip the [!TIP] download banner

    for raw_line in lines:
        line = raw_line.rstrip("\n")

        # ── Skip download tip block ──────────────────────────────────────
        if "Download" in line and ".docx" in line:
            skip_next = True
            continue
        if skip_next and line.startswith(">"):
            continue
        else:
            skip_next = False

        # ── Code block detection ─────────────────────────────────────────
        if line.strip().startswith("```"):
            if not in_code_block:
                in_code_block = True
                code_lines = []
            else:
                in_code_block = False
                add_code_block(doc, code_lines)
                code_lines = []
            continue

        if in_code_block:
            code_lines.append(line)
            continue

        # ── Table detection ──────────────────────────────────────────────
        if line.startswith("|"):
            row = [c.strip() for c in line.split("|") if c.strip()]
            # Skip separator rows like |---|---|
            if not all(re.match(r"^[-: ]+$", c) for c in row):
                table_rows.append(row)
            in_table = True
            continue
        elif in_table:
            add_table(doc, table_rows)
            table_rows = []
            in_table = False

        # ── Horizontal rule ──────────────────────────────────────────────
        if line.strip().startswith("---"):
            add_horizontal_rule(doc)
            continue

        # ── Blockquotes ──────────────────────────────────────────────────
        if line.startswith("> "):
            inner = line[2:].strip()
            alert_match = re.match(r"\[!(TIP|NOTE|IMPORTANT|CAUTION|WARNING)\]", inner)
            if alert_match:
                alert_type = alert_match.group(1)
                body = inner[alert_match.end():].strip()
                add_blockquote(doc, body, alert_type)
            else:
                add_blockquote(doc, inner)
            continue

        # ── Headings ─────────────────────────────────────────────────────
        heading_match = re.match(r"^(#{1,5}) (.+)", line)
        if heading_match:
            hashes = heading_match.group(1)
            title  = heading_match.group(2)
            level  = len(hashes) - 1  # # → 0, ## → 1, ### → 2
            add_styled_heading(doc, title, level)
            continue

        # ── Empty line ───────────────────────────────────────────────────
        if not line.strip():
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            continue

        # ── Regular body text ─────────────────────────────────────────────
        add_body_paragraph(doc, line)

    # Handle table at end of file
    if in_table and table_rows:
        add_table(doc, table_rows)

    print(f"Saving to {docx_path}...")
    doc.save(docx_path)
    print("[Done] Document saved successfully.")


if __name__ == "__main__":
    md_file   = "PROJECT_DOCUMENTATION.md"
    docx_file = "Federated_SSL_Project_Study_Notes.docx"
    create_docx(md_file, docx_file)
